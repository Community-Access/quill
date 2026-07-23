"""Native Word (.docx) reader into :class:`RichDocument` (Phase 7, docx rich mode).

The exact inverse of :mod:`quill.io.docx_writer`: it reads the same vocabulary
that writer emits — Heading 1-6 / List Bullet / Quote-Title-Subtitle-Caption
paragraph styles, per-run bold/italic/underline/strike/super/subscript, font
family and point size, text color, highlight, and paragraph alignment — into
the shared :class:`RichDocument` model. Reader and writer sharing one model is
what makes the docx rich round trip symmetric by construction.

Editable Word documents are *reconstructive*: no editor control hosts .docx
natively, so a rich edit session is docx -> RichDocument -> RTF -> the native
rich control, and save reverses it through ``rich_to_docx_bytes``. Anything
outside the vocabulary is therefore lost on save — which is why
:func:`scan_docx_features` inventories those features honestly *before* the
user commits, and why the UI backs up a flagged original before the first
overwrite. QUILL never silently rewrites someone's Word file.

python-docx is optional (as in the writer): :func:`python_docx_available`
gates; absent, callers stay on the classic read-extract path.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from quill.io.docx_writer import python_docx_available
from quill.io.rtf_model import InlineSpan, RichDocument, RichParagraph

__all__ = ["python_docx_available", "read_docx_rich", "scan_docx_features"]

#: Word built-in style names -> RichDocument named styles (writer table, inverted).
_NAMED_STYLES = {
    "Quote": "quote",
    "Title": "title",
    "Subtitle": "subtitle",
    "Caption": "caption",
}

#: WD_COLOR_INDEX member names -> the hidden-codes highlight vocabulary
#: (docx_writer._HIGHLIGHT_NAMES, inverted; GRAY_25 folds to "gray").
_HIGHLIGHT_VALUES = {
    "YELLOW": "yellow",
    "BRIGHT_GREEN": "green",
    "TURQUOISE": "turquoise",
    "PINK": "pink",
    "BLUE": "blue",
    "RED": "red",
    "GRAY_25": "gray",
}

_ALIGNMENTS = {"LEFT": "left", "RIGHT": "right", "CENTER": "center", "JUSTIFY": "justify"}


def _paragraph_shape(style_name: str) -> tuple[str, int, str | None]:
    """Map a Word style name to (style, heading level, named_style)."""
    name = (style_name or "").strip()
    if name.startswith("Heading "):
        try:
            level = int(name.removeprefix("Heading ").strip())
        except ValueError:
            level = 0
        if 1 <= level <= 6:
            return "heading", level, None
    if name.startswith("List Bullet"):
        return "bullet", 0, None
    if name.startswith("List Number"):
        return "numbered", 0, None
    if name in _NAMED_STYLES:
        return "paragraph", 0, _NAMED_STYLES[name]
    return "paragraph", 0, None


def _span_from_run(run: Any, *, href: str | None = None) -> InlineSpan:
    font = run.font
    color = None
    rgb = getattr(getattr(font, "color", None), "rgb", None)
    if rgb is not None:
        color = f"#{rgb!s}".lower() if not str(rgb).startswith("#") else str(rgb).lower()
    highlight = None
    highlight_value = getattr(font, "highlight_color", None)
    if highlight_value is not None:
        highlight = _HIGHLIGHT_VALUES.get(getattr(highlight_value, "name", str(highlight_value)))
    size = getattr(font, "size", None)
    return InlineSpan(
        text=str(run.text),
        bold=bool(run.bold),
        italic=bool(run.italic),
        underline=bool(run.underline),
        strike=bool(getattr(font, "strike", False)),
        superscript=bool(getattr(font, "superscript", False)),
        subscript=bool(getattr(font, "subscript", False)),
        font_family=getattr(font, "name", None),
        font_size_pt=int(size.pt) if size is not None else None,
        color=color,
        highlight=highlight,
        href=href,
    )


def _iter_body_blocks(source: Any) -> Any:
    """Yield the document body's paragraphs and tables in reading order.

    python-docx exposes ``source.paragraphs`` and ``source.tables`` as separate
    flat lists, losing their interleaving. Walking the body XML preserves order,
    so a table lands where it actually sits between paragraphs.
    """
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = source.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, source)
        elif isinstance(child, CT_Tbl):
            yield Table(child, source)


def _cell_text(cell: Any) -> str:
    """Flatten a table cell to a single line, escaped for a Markdown table."""
    text = " ".join(p.text for p in cell.paragraphs).strip()
    return text.replace("\\", "\\\\").replace("|", "\\|").replace("\n", " ")


def _table_to_rich_paragraphs(table: Any) -> list[RichParagraph]:
    """Render a Word table as GFM Markdown-table paragraphs.

    Rich edit has no native table vocabulary, so tables were previously dropped
    entirely -- their content vanished on open. Representing them as a Markdown
    table preserves the content, keeps it readable, and makes it reachable by
    single-key table navigation (the ``T`` quick-nav), while ``scan_docx_features``
    still flags the format change honestly.
    """
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([_cell_text(cell) for cell in row.cells])
    if not rows:
        return []
    columns = max(len(row) for row in rows)

    def _md_row(cells: list[str]) -> str:
        padded = cells + [""] * (columns - len(cells))
        return "| " + " | ".join(padded) + " |"

    lines = [_md_row(rows[0]), "| " + " | ".join(["---"] * columns) + " |"]
    lines.extend(_md_row(row) for row in rows[1:])
    return [RichParagraph(spans=[InlineSpan(text=line)]) for line in lines]


def _paragraph_spans(paragraph: Any) -> list[InlineSpan]:
    """Runs of a paragraph in reading order, carrying hyperlink hrefs.

    Walks ``iter_inner_content`` (Runs and Hyperlinks in document order) so a
    hyperlink's runs recover their target URL onto the span -- ``paragraph.runs``
    alone omits runs nested inside ``<w:hyperlink>``, which is why link hrefs
    were previously lost on a docx read.
    """
    from docx.text.hyperlink import Hyperlink
    from docx.text.run import Run

    spans: list[InlineSpan] = []
    iter_inner = getattr(paragraph, "iter_inner_content", None)
    if not callable(iter_inner):  # pragma: no cover - very old python-docx
        return [_span_from_run(run) for run in paragraph.runs if run.text]
    for item in iter_inner():
        if isinstance(item, Hyperlink):
            for run in item.runs:
                if run.text:
                    spans.append(_span_from_run(run, href=item.address or None))
        elif isinstance(item, Run) and item.text:
            spans.append(_span_from_run(item))
    return spans


def _paragraph_to_rich(paragraph: Any) -> RichParagraph:
    style, level, named_style = _paragraph_shape(getattr(paragraph.style, "name", ""))
    alignment = paragraph.alignment
    align = _ALIGNMENTS.get(getattr(alignment, "name", str(alignment))) if alignment else None
    spans = _paragraph_spans(paragraph)
    if not spans and paragraph.text:
        spans = [InlineSpan(text=str(paragraph.text))]
    return RichParagraph(
        spans=spans, style=style, level=level, align=align, named_style=named_style
    )


def read_docx_rich(path: Path) -> RichDocument:
    """Read ``path`` into a :class:`RichDocument` (the writer's vocabulary only).

    Raises ``ModuleNotFoundError`` when python-docx is absent (gate on
    :func:`python_docx_available` first) and ``OSError``/``ValueError`` for an
    unreadable file, matching the io-layer contract.

    Tables are rendered inline as Markdown-table paragraphs so their content is
    preserved and navigable (they were previously dropped entirely).
    """
    import docx
    from docx.table import Table

    source = docx.Document(str(path))
    rich = RichDocument()
    # Word's List Number style auto-numbers; the per-item number is not stored,
    # so assign 1, 2, 3... to each run of consecutive numbered paragraphs (the
    # count resets on any non-numbered paragraph) so the ordered list survives.
    numbered_counter = 0
    for block in _iter_body_blocks(source):
        if isinstance(block, Table):
            rich.paragraphs.extend(_table_to_rich_paragraphs(block))
            numbered_counter = 0
            continue
        para = _paragraph_to_rich(block)
        if para.style == "numbered":
            numbered_counter += 1
            para.list_number = numbered_counter
        else:
            numbered_counter = 0
        rich.paragraphs.append(para)
    return rich


def scan_docx_features(path: Path) -> list[str]:
    """Inventory the out-of-vocabulary features a rich edit would drop on save.

    The honest-fidelity gate for docx rich mode (the docx mirror of
    ``scan_rtf_features``): tables, images, comments, tracked changes,
    headers/footers, footnotes/endnotes, and extra sections are named up
    front so the user chooses — rich edit with these listed losses, the
    classic read-extract, or editing a copy. Best-effort and read-only: an
    unreadable package returns an empty list (the open path will surface the
    real error).
    """
    findings: list[str] = []
    try:
        import docx

        source = docx.Document(str(path))
    except Exception:  # noqa: BLE001 - the scan must never block the open path
        return findings
    try:
        if source.tables:
            findings.append(f"tables ({len(source.tables)})")
        if source.inline_shapes:
            findings.append(f"images ({len(source.inline_shapes)})")
        if len(source.sections) > 1:
            findings.append(f"multiple sections ({len(source.sections)})")
        for section in source.sections:
            header = section.header
            footer = section.footer
            if (
                header is not None
                and not header.is_linked_to_previous
                and any(p.text.strip() for p in header.paragraphs)
            ):
                findings.append("headers")
                break
        for section in source.sections:
            footer = section.footer
            if (
                footer is not None
                and not footer.is_linked_to_previous
                and any(p.text.strip() for p in footer.paragraphs)
            ):
                findings.append("footers")
                break
        body_xml = source.element.body.xml
        if "<w:ins " in body_xml or "<w:del " in body_xml:
            findings.append("tracked changes")
        if "<w:commentReference" in body_xml:
            findings.append("comments")
        if "<w:footnoteReference" in body_xml:
            findings.append("footnotes")
        if "<w:endnoteReference" in body_xml:
            findings.append("endnotes")
    except Exception:  # noqa: BLE001 - partial inventory beats a blocked open
        pass
    return findings
