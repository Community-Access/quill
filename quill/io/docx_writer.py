"""Native Word (.docx) writer built on python-docx.

The legacy export path runs Markdown through Pandoc, whose docx writer silently
drops the per-run font family, point size, color, highlight and per-paragraph
alignment that QUILL's hidden-codes feature
(``the QUILL PRD hidden-codes appendix``) materializes. This module
consumes the shared :class:`~quill.io.rtf_model.RichDocument` and writes those
attributes onto real Word runs and paragraphs, so a styled QUILL document opens in
Word with its formatting intact and its structure screen-reader navigable.

python-docx is an optional dependency. :func:`python_docx_available` reports
whether it is importable; callers (``quill.io.export.write_docx_document``) fall
back to the Pandoc path when it is not, so docx export never hard-fails.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING, Any

from quill.io.docx_math import omml_fragment_for_latex, split_math_segments
from quill.io.rtf_model import RichDocument, RichParagraph, markdown_to_rich

if TYPE_CHECKING:
    from quill.core.document import Document

__all__ = ["python_docx_available", "write_docx", "rich_to_docx_bytes"]

_NAMED_COLORS: dict[str, tuple[int, int, int]] = {
    "red": (255, 0, 0),
    "green": (0, 128, 0),
    "blue": (0, 0, 255),
    "black": (0, 0, 0),
    "white": (255, 255, 255),
    "yellow": (255, 255, 0),
    "orange": (255, 165, 0),
    "purple": (128, 0, 128),
    "gray": (128, 128, 128),
    "grey": (128, 128, 128),
}

#: python-docx highlight is an enum of named indices, not arbitrary RGB. Map the
#: common highlight names; anything else falls back to yellow.
_HIGHLIGHT_NAMES = {
    "yellow": "YELLOW",
    "green": "BRIGHT_GREEN",
    "bright_green": "BRIGHT_GREEN",
    "turquoise": "TURQUOISE",
    "pink": "PINK",
    "blue": "BLUE",
    "red": "RED",
    "gray": "GRAY_25",
    "grey": "GRAY_25",
}


def python_docx_available() -> bool:
    """Return ``True`` when python-docx can be imported."""
    try:
        import docx  # noqa: F401
    except ImportError:
        return False
    return True


def _parse_rgb(value: str) -> tuple[int, int, int] | None:
    text = value.strip()
    if text.startswith("#"):
        digits = text[1:]
        if len(digits) == 3:
            digits = "".join(char * 2 for char in digits)
        if len(digits) == 6:
            try:
                return (int(digits[0:2], 16), int(digits[2:4], 16), int(digits[4:6], 16))
            except ValueError:
                return None
        return None
    return _NAMED_COLORS.get(text.lower())


#: Named paragraph styles -> Word built-in style names.
_NAMED_STYLES = {
    "quote": "Quote",
    "title": "Title",
    "subtitle": "Subtitle",
    "caption": "Caption",
}
#: line-spacing attribute value -> python-docx multiple.
_LINE_SPACING = {"1": 1.0, "1.5": 1.5, "2": 2.0}


def _paragraph_style(paragraph: RichParagraph) -> str | None:
    if paragraph.style == "heading":
        level = min(max(paragraph.level, 1), 6)
        return f"Heading {level}"
    if paragraph.style == "bullet":
        return "List Bullet"
    if paragraph.style == "numbered":
        return "List Number"
    if paragraph.named_style in _NAMED_STYLES:
        return _NAMED_STYLES[paragraph.named_style]
    return None


def rich_to_docx(document: RichDocument) -> Any:
    """Build a python-docx ``Document`` from a :class:`RichDocument`.

    Raises ``ModuleNotFoundError`` (via the import) when python-docx is absent;
    callers should gate on :func:`python_docx_available` first.
    """
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.shared import Pt, RGBColor

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    out = docx.Document()

    def add_paragraph(paragraph: RichParagraph) -> None:
        para = out.add_paragraph()
        style = _paragraph_style(paragraph)
        if style is not None:
            try:
                para.style = style
            except KeyError:  # pragma: no cover - template without the named style
                pass
        if paragraph.align in align_map:
            para.alignment = align_map[paragraph.align]
        fmt = para.paragraph_format
        if paragraph.line_spacing in _LINE_SPACING:
            fmt.line_spacing = _LINE_SPACING[paragraph.line_spacing]
        if paragraph.space_before:
            fmt.space_before = Pt(paragraph.space_before)
        if paragraph.space_after:
            fmt.space_after = Pt(paragraph.space_after)
        if paragraph.indent:
            fmt.left_indent = Pt(paragraph.indent)
        if paragraph.first_line_indent:
            fmt.first_line_indent = Pt(paragraph.first_line_indent)
        for span in paragraph.spans:
            span_runs = []
            for segment in split_math_segments(span.text):
                if segment.is_math:
                    fragment = omml_fragment_for_latex(segment.content, display=segment.display)
                    if fragment is not None:
                        from docx.oxml import parse_xml

                        para._p.append(parse_xml(fragment))
                        continue
                    # Pandoc unavailable or conversion failed: keep the literal text
                    # rather than silently dropping the equation.
                    delimited = (
                        f"$${segment.content}$$" if segment.display else f"\\({segment.content}\\)"
                    )
                    run = para.add_run(delimited)
                else:
                    run = para.add_run(segment.content)
                span_runs.append(run)
                run.bold = span.bold or None
                run.italic = span.italic or None
                run.underline = span.underline or None
                if span.strike:
                    run.font.strike = True
                # superscript and subscript share one w:vertAlign element, so set only
                # the active one; assigning the other (even None) would clear it.
                if span.superscript:
                    run.font.superscript = True
                elif span.subscript:
                    run.font.subscript = True
                if span.font_family:
                    run.font.name = span.font_family
                if span.font_size_pt:
                    run.font.size = Pt(span.font_size_pt)
                if span.color:
                    rgb = _parse_rgb(span.color)
                    if rgb is not None:
                        run.font.color.rgb = RGBColor(*rgb)
                if span.highlight:
                    name = _HIGHLIGHT_NAMES.get(span.highlight.lower(), "YELLOW")
                    run.font.highlight_color = getattr(WD_COLOR_INDEX, name)
            if span.href and span_runs:
                _wrap_runs_in_hyperlink(para, span.href, span_runs)

    paras = document.paragraphs
    index = 0
    total = len(paras)
    while index < total:
        block = _table_block_length(paras, index)
        if block:
            _emit_table(out, paras[index : index + block])
            index += block
            # Two adjacent tables would merge into one in Word; separate them.
            if index < total and _table_block_length(paras, index):
                out.add_paragraph()
            continue
        paragraph = paras[index]
        index += 1
        if paragraph.style == "pagebreak":
            out.add_page_break()
            continue
        add_paragraph(paragraph)
    return out


_SEPARATOR_CELL_RE = re.compile(r"^:?-+:?$")


def _split_table_cells(line: str) -> list[str]:
    """Split a GFM table row (``| a | b |``) into unescaped cell strings.

    Mirrors :func:`quill.io.docx_reader._cell_text`'s escaping: cells escape a
    literal pipe as ``\\|`` and a backslash as ``\\\\``, so we split on pipes not
    preceded by a backslash and then unescape.
    """
    inner = line.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    cells = re.split(r"(?<!\\)\|", inner)
    return [c.strip().replace("\\|", "|").replace("\\\\", "\\") for c in cells]


def _is_plain_body(paragraph: RichParagraph) -> bool:
    """A default body paragraph (no heading/list/pagebreak style), the only kind
    a GFM table row is ever represented as."""
    return paragraph.style in ("", "paragraph", None)


def _is_table_row(text: str) -> bool:
    stripped = text.strip()
    return stripped.startswith("|") and stripped.endswith("|") and len(stripped) >= 2


def _is_separator_row(text: str) -> bool:
    if not _is_table_row(text):
        return False
    cells = _split_table_cells(text)
    return bool(cells) and all(_SEPARATOR_CELL_RE.match(c) for c in cells)


def _table_block_length(paras: list[RichParagraph], start: int) -> int:
    """Number of paragraphs forming a GFM table starting at ``start`` (0 if none).

    A table is a header row, a ``---`` separator row, then zero or more body
    rows -- exactly the shape :func:`docx_reader._table_to_rich_paragraphs`
    emits and that ``markdown_to_rich`` produces from a Markdown/HTML table.
    """
    if start + 1 >= len(paras):
        return 0
    if not _is_plain_body(paras[start]) or not _is_table_row(paras[start].text()):
        return 0
    if not _is_separator_row(paras[start + 1].text()):
        return 0
    length = 2
    while start + length < len(paras):
        row = paras[start + length]
        if not _is_plain_body(row) or not _is_table_row(row.text()):
            break
        if _is_separator_row(row.text()):
            break
        # A body row that is itself immediately followed by a separator is the
        # header of the *next* table -- stop so the tables stay distinct.
        following = paras[start + length + 1] if start + length + 1 < len(paras) else None
        if following is not None and _is_separator_row(following.text()):
            break
        length += 1
    return length


def _emit_table(out: Any, block: list[RichParagraph]) -> None:
    """Emit ``block`` (header, separator, body rows) as a real Word table.

    Reconstructs an editable ``w:tbl`` -- with a bold, repeated header row
    marked ``<w:tblHeader>`` so assistive technology announces column headers --
    instead of leaving the Markdown pipe-text as literal paragraphs.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    header = _split_table_cells(block[0].text())
    body = [_split_table_cells(p.text()) for p in block[2:]]
    columns = max([len(header)] + [len(r) for r in body] + [1])

    table = out.add_table(rows=1 + len(body), cols=columns)
    try:
        table.style = "Table Grid"
    except KeyError:  # pragma: no cover - template without the built-in style
        pass

    def _fill(cells: Any, values: list[str], *, bold: bool) -> None:
        for col, value in enumerate(values[:columns]):
            paragraph = cells[col].paragraphs[0]
            run = paragraph.add_run(value)
            run.bold = bold or None

    header_row = table.rows[0]
    _fill(header_row.cells, header, bold=True)
    # Mark the first row as a repeating header (accessibility + long tables).
    trPr = header_row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)

    for offset, values in enumerate(body, start=1):
        _fill(table.rows[offset].cells, values, bold=False)


def _wrap_runs_in_hyperlink(paragraph: Any, url: str, runs: list) -> None:
    """Move ``runs`` into a real ``<w:hyperlink>`` with an external relationship.

    python-docx has no high-level hyperlink writer, so this is the OPC recipe:
    relate the URL, create the ``w:hyperlink`` element in the run's place, and
    reparent the run(s) into it -- so the href round-trips (read back via
    ``iter_inner_content``) instead of a link degrading to plain text.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    runs[0]._r.addprevious(hyperlink)
    for run in runs:
        hyperlink.append(run._r)


def rich_to_docx_bytes(document: RichDocument) -> bytes:
    """Serialize a :class:`RichDocument` to in-memory ``.docx`` bytes."""
    import io

    buffer = io.BytesIO()
    rich_to_docx(document).save(buffer)
    return buffer.getvalue()


def write_docx(document: Document, target: Path) -> Path:
    """Write ``document``'s canonical markup to ``target`` as a native ``.docx``.

    The markup is converted to the shared :class:`RichDocument` first, so headings,
    bullets, emphasis, and all hidden-codes attributes (font, size, color,
    highlight, alignment) are carried onto Word runs and paragraphs. When the
    document has a Header/Footer Builder spec (#892), it is written as real
    Word header/footer parts with a live PAGE field.
    """
    import io

    from quill.core.storage import write_bytes_atomic

    rich = markdown_to_rich(document.text)
    built = rich_to_docx(rich)
    _maybe_apply_header_footer(built, document, target)
    # Atomic write: serialize to memory, then write the bytes via temp +
    # os.replace so a failure mid-serialize never destroys the previous file at
    # ``target``.
    buffer = io.BytesIO()
    built.save(buffer)
    write_bytes_atomic(target, buffer.getvalue())
    return target


def _maybe_apply_header_footer(built: Any, document: Document, target: Path) -> None:
    """Apply the document's saved Header/Footer spec to the built docx, if any.

    Best-effort by contract: the header/footer is presentation, and its export
    must never be the reason a save fails.
    """
    try:
        import datetime

        from quill.core.header_footer_store import HeaderFooterStore, key_for
        from quill.io.header_footer_export import docx_apply_header_footer, spec_has_content

        spec = HeaderFooterStore.load().get(key_for(document.path or target))
        if not spec_has_content(spec) or spec is None:
            return
        name = Path(target).name
        docx_apply_header_footer(
            built,
            spec,
            title=name.rsplit(".", 1)[0] if "." in name else name,
            filename=name,
            date=datetime.date.today().isoformat(),
        )
    except Exception:  # noqa: BLE001 - never fail the save over a header
        return
