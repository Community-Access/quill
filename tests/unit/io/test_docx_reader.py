"""docx_reader: the inverse of docx_writer through the shared RichDocument.

The docx rich round trip is reconstructive, so its correctness rests on one
property: **reader and writer share the RichDocument vocabulary exactly**.
These tests write a document through ``rich_to_docx_bytes`` and read it back
with ``read_docx_rich``, asserting identity across the full vocabulary, plus
the ``scan_docx_features`` honest-fidelity inventory.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from quill.io.docx_reader import python_docx_available, read_docx_rich, scan_docx_features
from quill.io.docx_writer import rich_to_docx_bytes
from quill.io.rtf_model import InlineSpan, RichDocument, RichParagraph

pytestmark = pytest.mark.skipif(not python_docx_available(), reason="python-docx not installed")


def _write(tmp_path: Path, document: RichDocument, name: str = "doc.docx") -> Path:
    target = tmp_path / name
    target.write_bytes(rich_to_docx_bytes(document))
    return target


def test_vocabulary_round_trips_docx_to_rich_identity(tmp_path: Path) -> None:
    source = RichDocument(
        paragraphs=[
            RichParagraph(spans=[InlineSpan(text="Title line")], style="heading", level=1),
            RichParagraph(
                spans=[
                    InlineSpan(text="plain "),
                    InlineSpan(text="bold", bold=True),
                    InlineSpan(text=" and "),
                    InlineSpan(text="italic", italic=True),
                    InlineSpan(text=" and "),
                    InlineSpan(text="underline", underline=True),
                ]
            ),
            RichParagraph(
                spans=[
                    InlineSpan(
                        text="styled",
                        font_family="Consolas",
                        font_size_pt=14,
                        color="#ff0000",
                        highlight="yellow",
                    )
                ],
                align="center",
            ),
            RichParagraph(spans=[InlineSpan(text="a bullet")], style="bullet"),
            RichParagraph(
                spans=[
                    InlineSpan(text="struck", strike=True),
                    InlineSpan(text="sup", superscript=True),
                ]
            ),
        ]
    )
    loaded = read_docx_rich(_write(tmp_path, source))

    headings = [p for p in loaded.paragraphs if p.style == "heading"]
    assert len(headings) == 1 and headings[0].level == 1
    assert headings[0].text() == "Title line"

    emphasis = loaded.paragraphs[1]
    assert emphasis.text() == "plain bold and italic and underline"
    flags = {(s.text, s.bold, s.italic, s.underline) for s in emphasis.spans}
    assert ("bold", True, False, False) in flags
    assert ("italic", False, True, False) in flags
    assert ("underline", False, False, True) in flags

    styled = loaded.paragraphs[2]
    assert styled.align == "center"
    span = styled.spans[0]
    assert span.font_family == "Consolas"
    assert span.font_size_pt == 14
    assert span.color == "#ff0000"
    assert span.highlight == "yellow"

    assert any(p.style == "bullet" and p.text() == "a bullet" for p in loaded.paragraphs)
    last = loaded.paragraphs[4]
    assert any(s.strike for s in last.spans)
    assert any(s.superscript for s in last.spans)


def test_named_styles_round_trip(tmp_path: Path) -> None:
    source = RichDocument(
        paragraphs=[RichParagraph(spans=[InlineSpan(text="quoted")], named_style="quote")]
    )
    loaded = read_docx_rich(_write(tmp_path, source))
    assert loaded.paragraphs[0].named_style == "quote"


def test_scan_flags_tables_and_images(tmp_path: Path) -> None:
    import docx

    document = docx.Document()
    document.add_paragraph("hello")
    document.add_table(rows=2, cols=2)
    target = tmp_path / "flagged.docx"
    document.save(str(target))
    findings = scan_docx_features(target)
    assert any(f.startswith("tables") for f in findings)


def test_scan_is_quiet_for_a_clean_document(tmp_path: Path) -> None:
    source = RichDocument(paragraphs=[RichParagraph(spans=[InlineSpan(text="clean")])])
    assert scan_docx_features(_write(tmp_path, source)) == []


def test_scan_never_raises_on_junk(tmp_path: Path) -> None:
    junk = tmp_path / "junk.docx"
    junk.write_bytes(b"this is not a zip")
    assert scan_docx_features(junk) == []


def test_tables_are_preserved_inline_as_markdown_not_dropped(tmp_path: Path) -> None:
    # A Word table used to vanish entirely on rich read (its cells were lost).
    # It now appears inline, in document order, as a GFM Markdown table so the
    # content survives and is reachable by single-key table navigation.
    import docx

    source_doc = docx.Document()
    source_doc.add_paragraph("Intro paragraph.")
    table = source_doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "First Name"
    table.cell(0, 1).text = "Last Name"
    table.cell(1, 0).text = "Ada"
    table.cell(1, 1).text = "Lovelace"
    source_doc.add_paragraph("After the table.")
    path = tmp_path / "with_table.docx"
    source_doc.save(str(path))

    rich = read_docx_rich(path)
    lines = ["".join(span.text for span in par.spans) for par in rich.paragraphs]
    assert lines == [
        "Intro paragraph.",
        "| First Name | Last Name |",
        "| --- | --- |",
        "| Ada | Lovelace |",
        "After the table.",
    ]


def test_table_cell_pipes_are_escaped(tmp_path: Path) -> None:
    import docx

    source_doc = docx.Document()
    table = source_doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "a | b"
    path = tmp_path / "pipe.docx"
    source_doc.save(str(path))

    rich = read_docx_rich(path)
    first = "".join(span.text for span in rich.paragraphs[0].spans)
    assert first == r"| a \| b |"


def test_imported_table_navigates_after_the_rich_edit_rtf_round_trip(tmp_path: Path) -> None:
    """A Word table stays cell-navigable through the whole rich-edit path.

    Rich Word mode is docx -> RichDocument -> RTF -> the native control, so the
    text the caret actually sees is the far side of that round trip. The escaped
    pipe a cell needs has to survive it, or the grid comes back one column wide
    and every announced position is wrong.
    """
    import docx

    from quill.core import table_nav
    from quill.io.rtf_model import rich_to_rtf, rtf_to_rich

    source_doc = docx.Document()
    table = source_doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Size | Units"
    table.cell(1, 0).text = "Depth"
    table.cell(1, 1).text = "3"
    path = tmp_path / "round_trip.docx"
    source_doc.save(str(path))

    landed = rtf_to_rich(rich_to_rtf(read_docx_rich(path)))
    text = "\n".join("".join(s.text for s in par.spans) for par in landed.paragraphs) + "\n"
    grid = table_nav.find_table_at(text, text.index("Depth"))

    assert grid is not None
    assert (grid.row_count, grid.col_count) == (2, 2)
    assert table_nav.move(grid, text.index("Depth"), "up").announcement == (
        "Row 1 of 2, column 1 of 2: Metric"
    )
    assert table_nav.move(grid, text.index("Metric"), "next").announcement == (
        "Row 1 of 2, column 2 of 2: Size | Units"
    )
