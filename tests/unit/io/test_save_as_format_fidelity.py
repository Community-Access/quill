"""Save As fidelity across formats: headings, lists, bold, italic, underline,
and links must come across into each target's NATIVE markup.

The scenario that motivated this: open a Word (.docx) document in rich mode and
Save As Markdown/HTML. In rich mode the editor's ``document.text`` is a flattened
plain-text mirror, so the save path must convert the *RichDocument* (reconstructed
from the control's RTF) rather than serialize that mirror. These tests exercise
the pure io-layer converters that back that path, plus the Markdown-source path
used when a document is opened on the extract floor.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from quill.io.export import markdown_to_html
from quill.io.rtf_model import (
    InlineSpan,
    RichDocument,
    RichParagraph,
    markdown_to_rich,
    rich_to_markdown,
)


def _word_like_document() -> RichDocument:
    """A RichDocument shaped like a Word doc opened in rich mode: a heading, a
    paragraph mixing bold/italic/underline/link, and a two-item bullet list."""
    return RichDocument(
        paragraphs=[
            RichParagraph(spans=[InlineSpan(text="Report Title")], style="heading", level=1),
            RichParagraph(
                spans=[
                    InlineSpan(text="normal "),
                    InlineSpan(text="bold", bold=True),
                    InlineSpan(text=" "),
                    InlineSpan(text="italic", italic=True),
                    InlineSpan(text=" "),
                    InlineSpan(text="under", underline=True),
                    InlineSpan(text=" "),
                    InlineSpan(text="the site", href="https://example.com"),
                ]
            ),
            RichParagraph(spans=[InlineSpan(text="First point")], style="bullet"),
            RichParagraph(spans=[InlineSpan(text="Second point")], style="bullet"),
        ]
    )


# --------------------------------------------------------------------------- #
# Rich (opened-from-Word) -> Markdown
# --------------------------------------------------------------------------- #


def test_rich_to_markdown_uses_native_markdown() -> None:
    md = rich_to_markdown(_word_like_document())
    assert "# Report Title" in md  # heading
    assert "**bold**" in md  # bold
    assert "*italic*" in md  # italic
    assert "[the site](https://example.com)" in md  # link
    assert "- First point" in md and "- Second point" in md  # bullet list


def test_rich_to_markdown_keeps_underline_via_quill_span() -> None:
    # Markdown has no standard underline syntax; QUILL preserves it as a
    # hidden-codes span that round-trips and renders in QUILL's preview/HTML.
    md = rich_to_markdown(_word_like_document())
    assert "[under]{underline}" in md


# --------------------------------------------------------------------------- #
# Rich (opened-from-Word) -> HTML  (everything, including underline, native)
# --------------------------------------------------------------------------- #


def test_rich_to_html_is_fully_native() -> None:
    html = markdown_to_html(rich_to_markdown(_word_like_document()), "Report")
    assert "<h1" in html and ">Report Title</h1>" in html
    assert "<strong>bold</strong>" in html
    assert "<em>italic</em>" in html
    assert '<a href="https://example.com">the site</a>' in html
    assert "<li>First point</li>" in html and "<li>Second point</li>" in html
    # Underline is native in HTML (unlike Markdown).
    assert "text-decoration: underline" in html
    assert ">under</span>" in html


# --------------------------------------------------------------------------- #
# Rich (opened-from-Word) -> DOCX  (round-trips through the writer/reader)
# --------------------------------------------------------------------------- #


def test_rich_to_docx_round_trips_formatting(tmp_path: Path) -> None:
    docx_writer = pytest.importorskip("quill.io.docx_writer")
    if not docx_writer.python_docx_available():
        pytest.skip("python-docx not installed")
    from quill.io.docx_reader import read_docx_rich

    data = docx_writer.rich_to_docx_bytes(_word_like_document())
    path = tmp_path / "out.docx"
    path.write_bytes(data)
    back = read_docx_rich(path)

    heading = back.paragraphs[0]
    assert heading.style == "heading" and heading.level == 1
    assert "Report Title" in "".join(s.text for s in heading.spans)

    body = back.paragraphs[1]
    flags = {(s.text.strip(), s.bold, s.italic, s.underline) for s in body.spans}
    assert ("bold", True, False, False) in flags
    assert ("italic", False, True, False) in flags
    assert ("under", False, False, True) in flags  # underline survives in docx
    # Hyperlinks now round-trip through docx (real <w:hyperlink> + relationship).
    assert any(s.href == "https://example.com" and "the site" in s.text for s in body.spans)

    bullets = [p for p in back.paragraphs if p.style == "bullet"]
    assert [s.text for p in bullets for s in p.spans] == ["First point", "Second point"]


def test_rich_to_docx_preserves_hyperlink_with_its_run_formatting(tmp_path: Path) -> None:
    docx_writer = pytest.importorskip("quill.io.docx_writer")
    if not docx_writer.python_docx_available():
        pytest.skip("python-docx not installed")
    from quill.io.docx_reader import read_docx_rich
    from quill.io.rtf_model import InlineSpan as _Span
    from quill.io.rtf_model import RichDocument as _Doc
    from quill.io.rtf_model import RichParagraph as _Para

    doc = _Doc(
        paragraphs=[_Para(spans=[_Span(text="a bold link", href="https://example.org", bold=True)])]
    )
    path = tmp_path / "link.docx"
    path.write_bytes(docx_writer.rich_to_docx_bytes(doc))
    span = read_docx_rich(path).paragraphs[0].spans[0]
    assert span.href == "https://example.org"
    assert span.bold is True
    assert span.text == "a bold link"


# --------------------------------------------------------------------------- #
# Markdown-source (extract floor) -> HTML / DOCX
# --------------------------------------------------------------------------- #

_MARKDOWN_SOURCE = "\n".join([
    "# Guide",
    "",
    "Intro with **bold**, *italic*, and [a link](https://quill.example).",
    "",
    "- alpha",
    "- beta",
])


def test_numbered_lists_round_trip_and_keep_their_start() -> None:
    # Markdown identity, including a non-1 start (the number is preserved).
    md = "1. First\n2. Second\n3. Third"
    assert rich_to_markdown(markdown_to_rich(md)) == md
    assert rich_to_markdown(markdown_to_rich("3. Three\n4. Four")) == "3. Three\n4. Four"


def test_numbered_list_to_html_is_native_ordered_list() -> None:
    html = markdown_to_html("1. one\n2. two", "t")
    assert "<ol" in html
    assert "<li>one</li>" in html and "<li>two</li>" in html


def test_numbered_list_round_trips_through_docx(tmp_path: Path) -> None:
    docx_writer = pytest.importorskip("quill.io.docx_writer")
    if not docx_writer.python_docx_available():
        pytest.skip("python-docx not installed")
    from quill.io.docx_reader import read_docx_rich

    doc = RichDocument(
        paragraphs=[
            RichParagraph(spans=[InlineSpan(text="Alpha")], style="numbered", list_number=1),
            RichParagraph(spans=[InlineSpan(text="Beta")], style="numbered", list_number=2),
        ]
    )
    path = tmp_path / "num.docx"
    path.write_bytes(docx_writer.rich_to_docx_bytes(doc))
    back = read_docx_rich(path)
    numbered = [(p.list_number, p.spans[0].text) for p in back.paragraphs if p.style == "numbered"]
    assert numbered == [(1, "Alpha"), (2, "Beta")]


def test_image_renders_as_img_in_html() -> None:
    # Regression: ![alt](src) used to render as a literal "!" plus a hyperlink.
    html = markdown_to_html("![a cat](cat.png)", "t")
    assert '<img src="cat.png" alt="a cat">' in html
    assert "!<a" not in html


def test_image_and_link_coexist_in_html() -> None:
    html = markdown_to_html("see ![pic](a.png) and [link](https://b.example)", "t")
    assert '<img src="a.png" alt="pic">' in html
    assert '<a href="https://b.example">link</a>' in html


def test_blockquote_renders_as_blockquote_in_html() -> None:
    # Regression: "> quote" used to stay literal (escaped &gt;) instead of a
    # real <blockquote> element.
    html = markdown_to_html("> quote one\n> quote two", "t")
    assert "<blockquote>quote one<br>quote two</blockquote>" in html


def test_gfm_table_to_docx_becomes_a_real_word_table(tmp_path: Path) -> None:
    # A Markdown/HTML table now becomes a real, editable Word table object --
    # not literal pipe-text paragraphs.
    docx = pytest.importorskip("docx")
    docx_writer = pytest.importorskip("quill.io.docx_writer")
    if not docx_writer.python_docx_available():
        pytest.skip("python-docx not installed")
    md = "| Name | Qty |\n| --- | --- |\n| Apples | 3 |\n| Pears | 12 |"
    path = tmp_path / "t.docx"
    path.write_bytes(docx_writer.rich_to_docx_bytes(markdown_to_rich(md)))
    doc = docx.Document(str(path))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert (len(table.rows), len(table.columns)) == (3, 2)  # header + 2 body rows
    grid = [[cell.text for cell in row.cells] for row in table.rows]
    assert grid == [["Name", "Qty"], ["Apples", "3"], ["Pears", "12"]]
    # No stray pipe-text paragraphs left behind.
    assert not any("|" in p.text for p in doc.paragraphs)


def test_gfm_table_to_docx_marks_a_repeating_header_row(tmp_path: Path) -> None:
    # The header row is a real repeating table header (<w:tblHeader>) so screen
    # readers announce column headers and Word repeats it across page breaks.
    pytest.importorskip("docx")
    docx_writer = pytest.importorskip("quill.io.docx_writer")
    if not docx_writer.python_docx_available():
        pytest.skip("python-docx not installed")
    import docx as _docx
    from docx.oxml.ns import qn

    md = "| H1 | H2 |\n| --- | --- |\n| a | b |"
    path = tmp_path / "hdr.docx"
    path.write_bytes(docx_writer.rich_to_docx_bytes(markdown_to_rich(md)))
    table = _docx.Document(str(path)).tables[0]
    header_tr = table.rows[0]._tr
    assert header_tr.find(qn("w:trPr")).find(qn("w:tblHeader")) is not None
    # Header cells are bold.
    assert table.rows[0].cells[0].paragraphs[0].runs[0].bold is True


def test_table_round_trips_markdown_to_word_and_back(tmp_path: Path) -> None:
    # Full loop: Markdown table -> real Word table -> read back -> Markdown table.
    pytest.importorskip("docx")
    docx_writer = pytest.importorskip("quill.io.docx_writer")
    if not docx_writer.python_docx_available():
        pytest.skip("python-docx not installed")
    from quill.io.docx_reader import read_docx_rich

    md = "| Fruit | Count |\n| --- | --- |\n| Apples | 3 |\n| Pears | 12 |"
    path = tmp_path / "rt.docx"
    path.write_bytes(docx_writer.rich_to_docx_bytes(markdown_to_rich(md)))
    assert rich_to_markdown(read_docx_rich(path)) == md


def test_word_table_survives_save_as_markdown_and_html() -> None:
    # A Word table (read as GFM paragraphs) exports as a native Markdown table
    # and a native HTML <table> -- the "Word table -> Markdown/HTML" direction.
    table_doc = RichDocument(
        paragraphs=[
            RichParagraph(spans=[InlineSpan(text="| City | Pop |")]),
            RichParagraph(spans=[InlineSpan(text="| --- | --- |")]),
            RichParagraph(spans=[InlineSpan(text="| Reno | 250k |")]),
        ]
    )
    md = rich_to_markdown(table_doc)
    assert md == "| City | Pop |\n| --- | --- |\n| Reno | 250k |"
    html = markdown_to_html(md, "t")
    assert "<table" in html and "<th>City</th>" in html and "<td>Reno</td>" in html


def test_table_with_surrounding_paragraphs_and_two_tables(tmp_path: Path) -> None:
    # Tables detected in a mixed document, and two adjacent tables stay distinct.
    docx = pytest.importorskip("docx")
    docx_writer = pytest.importorskip("quill.io.docx_writer")
    if not docx_writer.python_docx_available():
        pytest.skip("python-docx not installed")
    md = "\n".join([
        "# Title",
        "| A | B |",
        "| --- | --- |",
        "| 1 | 2 |",
        "| C | D |",
        "| --- | --- |",
        "| 3 | 4 |",
    ])
    path = tmp_path / "two.docx"
    path.write_bytes(docx_writer.rich_to_docx_bytes(markdown_to_rich(md)))
    doc = docx.Document(str(path))
    assert len(doc.tables) == 2
    assert [[c.text for c in r.cells] for r in doc.tables[0].rows] == [["A", "B"], ["1", "2"]]
    assert [[c.text for c in r.cells] for r in doc.tables[1].rows] == [["C", "D"], ["3", "4"]]


def test_markdown_source_to_html_is_native() -> None:
    html = markdown_to_html(_MARKDOWN_SOURCE, "Guide")
    assert "<h1" in html and ">Guide</h1>" in html
    assert "<strong>bold</strong>" in html
    assert "<em>italic</em>" in html
    assert '<a href="https://quill.example">a link</a>' in html
    assert "<li>alpha</li>" in html and "<li>beta</li>" in html


def test_markdown_source_to_docx_is_native(tmp_path: Path) -> None:
    docx_writer = pytest.importorskip("quill.io.docx_writer")
    if not docx_writer.python_docx_available():
        pytest.skip("python-docx not installed")
    from quill.io.docx_reader import read_docx_rich

    data = docx_writer.rich_to_docx_bytes(markdown_to_rich(_MARKDOWN_SOURCE))
    path = tmp_path / "guide.docx"
    path.write_bytes(data)
    back = read_docx_rich(path)

    assert back.paragraphs[0].style == "heading" and back.paragraphs[0].level == 1
    all_spans = [(s.text, s.bold, s.italic) for p in back.paragraphs for s in p.spans]
    assert ("bold", True, False) in all_spans
    assert ("italic", False, True) in all_spans
    bullets = [s.text for p in back.paragraphs if p.style == "bullet" for s in p.spans]
    assert bullets == ["alpha", "beta"]
